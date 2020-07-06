from PIL            import Image
import docx
from docx.shared    import RGBColor, Pt, Cm
from os             import chdir

chdir('C:/users/leona/Pictures/MonaLisas')

doc = docx.Document()

im = Image.open('mona.png')
name = im.filename
w, h = im.size
if w<1.5*h: 
    im = im.crop((0, 0, w, w//1.5))
    im.save(f'crop_{name}')
    w, h = im.size
amp  = 70/w
im = im.resize((int(amp*w), int(amp*h)))
w, h = im.size


pixels = []
for x in range(w):
    for y in range(h):
        pixels.append({(x, y):im.getpixel((x, y))})

paragraph_format = doc.styles['No Spacing'].paragraph_format

steps = h//20
#doc.add_heading('Mona Lisa', 0)
i = 0
for y in range(0, h, steps):
    para = doc.add_paragraph()
    
    for x in range(w):
        trecho = para.add_run('MonaLisa'[x%len('Monalisa')])
        r, g, b = im.getpixel((x, y))
        trecho.font.color.rgb = RGBColor(r, g, b)
    i += 1

doc.paragraphs[i-1].runs[w-1].add_break(docx.enum.text.WD_BREAK.PAGE)
amp = 10/max(w, h)
doc.add_picture(f'crop_{name}', width=docx.shared.Cm(w*amp), height=docx.shared.Cm(h*amp))
doc.save('mona.docx')
